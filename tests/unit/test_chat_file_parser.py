from io import BytesIO

from docx import Document
from openpyxl import Workbook

import pytest

from file_processing.parser import UnsupportedFileTypeError

from file_processing.parser import ChatFileParser


def test_parse_docx_text_and_table():
    doc = Document()
    doc.add_heading("实验报告", level=1)
    doc.add_paragraph("样品 A 的冲击强度为 35 kJ/m²。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "测试温度"
    table.rows[0].cells[1].text = "25 ℃"
    buf = BytesIO()
    doc.save(buf)

    parsed = ChatFileParser().parse_bytes("report.docx", buf.getvalue())

    assert parsed.parser == "python-docx"
    assert "35 kJ/m²" in parsed.text
    assert "测试温度" in parsed.text
    assert len(parsed.chunks) >= 1


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "配方"
    sheet.append(["样品", "ABS", "PC", "增韧剂", "冲击强度"])
    sheet.append(["ABS-051", 30, 50, 20, 35.2])
    sheet.append(["ABS-052", 35, 47, 18, 41.8])
    process = workbook.create_sheet("工艺")
    process.append(["样品", "加工温度", "螺杆转速"])
    process.append(["ABS-051", 250, 320])
    buf = BytesIO()
    workbook.save(buf)
    workbook.close()
    return buf.getvalue()


def test_parse_xlsx_preserves_sheet_rows_and_values():
    parsed = ChatFileParser(chunk_chars=400).parse_bytes(
        "实验数据.xlsx",
        _xlsx_bytes(),
    )

    assert parsed.parser == "openpyxl"
    assert parsed.media_type == (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert "[工作表] 配方" in parsed.text
    assert "ABS-051" in parsed.text
    assert "35.2" in parsed.text
    assert "[工作表] 工艺" in parsed.text
    assert any(chunk.sheet_name == "配方" for chunk in parsed.chunks)
    assert any(chunk.sheet_name == "工艺" for chunk in parsed.chunks)
    formula_chunk = next(chunk for chunk in parsed.chunks if chunk.sheet_name == "配方")
    assert formula_chunk.row_start == 1
    assert formula_chunk.row_end >= 3


def test_xls_remains_explicitly_unsupported():
    with pytest.raises(UnsupportedFileTypeError, match="PDF、DOCX 和 XLSX"):
        ChatFileParser().parse_bytes("legacy.xls", b"not-an-xls")


def test_parse_xlsx_reports_exact_sheet_metadata():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "工艺"
    # Simulate a wide/long sheet shape without making a huge fixture.
    for col in range(1, 72):
        ws.cell(row=1, column=col, value=f"字段{col}")
    for row in range(2, 96):
        ws.cell(row=row, column=1, value=f"S{row:03d}")
        ws.cell(row=row, column=2, value=row)
    buf = BytesIO()
    wb.save(buf)

    parsed = ChatFileParser(chunk_chars=1400).parse_bytes("工艺.xlsx", buf.getvalue())

    assert parsed.parser == "openpyxl"
    assert "原始最大行=95" in parsed.text
    assert "原始最大列=71" in parsed.text
    assert "非空行=95" in parsed.text
    assert parsed.chunks[-1].row_end == 95
    assert "原始最大行=95" in parsed.chunks[-1].text


def test_generic_xlsx_summary_selects_tail_chunks():
    from types import SimpleNamespace

    from skills.current_attachment import CurrentAttachmentSkill

    chunks = tuple(
        {
            "index": index,
            "text": f"[工作表] 工艺；[行范围] {index * 10 + 1}-{index * 10 + 10}\n内容{index}",
            "sheet_name": "工艺",
            "row_start": index * 10 + 1,
            "row_end": index * 10 + 10,
        }
        for index in range(13)
    )
    attachment = SimpleNamespace(
        attachment_id="xlsx-1",
        filename="工艺.xlsx",
        parser="openpyxl",
        chunks=chunks,
        chunk_count=len(chunks),
    )
    skill = CurrentAttachmentSkill(store=None, llm=None)  # type: ignore[arg-type]

    selected = skill._select_chunks("请分析总结这个附件", [attachment], limit=12)

    assert len(selected) == 13
    assert selected[-1]["chunk"]["index"] == 12
    assert skill._coverage_warnings(selected, [attachment], "请分析总结这个附件") == []
