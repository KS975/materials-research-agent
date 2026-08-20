from __future__ import annotations

from datetime import date, datetime, time
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from file_processing.models import ParsedChunk, ParsedDocument


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class UnifiedFileParser:
    """Unified PDF/DOCX/XLSX parser.

    The same parser output is reused by both:
    - current Chat temporary attachments
    - long-term Knowledge Index ingestion

    The parser itself has no persistence side effects: it writes neither
    ChatAttachmentStore nor Qdrant.
    """

    SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx"}
    XLSX_MAX_SHEETS = 20
    XLSX_MAX_ROWS_PER_SHEET = 5000
    XLSX_MAX_COLUMNS = 100

    def __init__(self, chunk_chars: int = 1400, overlap_chars: int = 180):
        self.chunk_chars = max(400, chunk_chars)
        self.overlap_chars = max(0, min(overlap_chars, self.chunk_chars // 2))

    def parse_bytes(self, filename: str, content: bytes, media_type: str = "") -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise UnsupportedFileTypeError("当前只支持 PDF、DOCX 和 XLSX")
        if not content:
            raise EmptyDocumentError("上传文件为空")

        if suffix == ".pdf":
            return self._parse_pdf(filename, content, media_type or "application/pdf")
        if suffix == ".xlsx":
            return self._parse_xlsx(
                filename,
                content,
                media_type
                or "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        return self._parse_docx(
            filename,
            content,
            media_type
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _parse_pdf(self, filename: str, content: bytes, media_type: str) -> ParsedDocument:
        reader = PdfReader(BytesIO(content))
        chunks: list[ParsedChunk] = []
        all_text: list[str] = []
        index = 0

        for page_number, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if not page_text:
                continue
            all_text.append(page_text)
            for piece in self._chunk_text(page_text):
                chunks.append(ParsedChunk(index=index, text=piece, page=page_number))
                index += 1

        text = "\n\n".join(all_text).strip()
        if not text:
            raise EmptyDocumentError(
                "PDF 未提取到可读文本。若这是扫描件，后续需要接 MinerU/OCR 解析器。"
            )

        return ParsedDocument(
            filename=filename,
            media_type=media_type,
            parser="pypdf",
            text=text,
            chunks=tuple(chunks),
            page_count=len(reader.pages),
        )

    def _parse_docx(self, filename: str, content: bytes, media_type: str) -> ParsedDocument:
        doc = DocxDocument(BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Include table cell text so experiment tables are not silently lost.
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                line = " | ".join(value for value in values if value)
                if line:
                    paragraphs.append(line)

        text = "\n\n".join(paragraphs).strip()
        if not text:
            raise EmptyDocumentError("DOCX 未提取到可读文本")

        chunks: list[ParsedChunk] = []
        buffer: list[str] = []
        buffer_len = 0
        start = 1
        index = 0

        def flush(end: int) -> None:
            nonlocal buffer, buffer_len, start, index
            if not buffer:
                return
            chunk_text = "\n\n".join(buffer).strip()
            chunks.append(
                ParsedChunk(
                    index=index,
                    text=chunk_text,
                    paragraph_start=start,
                    paragraph_end=end,
                )
            )
            index += 1
            buffer = []
            buffer_len = 0
            start = end + 1

        for paragraph_index, paragraph in enumerate(paragraphs, start=1):
            extra = len(paragraph) + (2 if buffer else 0)
            if buffer and buffer_len + extra > self.chunk_chars:
                flush(paragraph_index - 1)
            buffer.append(paragraph)
            buffer_len += extra
        flush(len(paragraphs))

        return ParsedDocument(
            filename=filename,
            media_type=media_type,
            parser="python-docx",
            text=text,
            chunks=tuple(chunks),
            page_count=None,
        )

    def _parse_xlsx(self, filename: str, content: bytes, media_type: str) -> ParsedDocument:
        try:
            workbook = load_workbook(
                BytesIO(content),
                read_only=True,
                data_only=False,
            )
        except Exception as exc:
            raise EmptyDocumentError(f"XLSX 无法读取：{type(exc).__name__}: {exc}") from exc

        chunks: list[ParsedChunk] = []
        document_sections: list[str] = []
        chunk_index = 0

        try:
            worksheets = workbook.worksheets[: self.XLSX_MAX_SHEETS]
            for worksheet in worksheets:
                max_column = min(
                    max(1, int(worksheet.max_column or 1)),
                    self.XLSX_MAX_COLUMNS,
                )
                rows: list[tuple[int, str]] = []
                for row_number, row in enumerate(
                    worksheet.iter_rows(
                        min_row=1,
                        max_row=self.XLSX_MAX_ROWS_PER_SHEET,
                        max_col=max_column,
                        values_only=True,
                    ),
                    start=1,
                ):
                    values = [self._xlsx_cell_text(value) for value in row]
                    # Trim trailing empty cells while preserving interior blanks.
                    while values and not values[-1]:
                        values.pop()
                    if not values or not any(values):
                        continue
                    line = " | ".join(values)
                    rows.append((row_number, line))

                if not rows:
                    continue

                truncation_notes: list[str] = []
                if int(worksheet.max_row or 0) > self.XLSX_MAX_ROWS_PER_SHEET:
                    truncation_notes.append(
                        f"仅解析前 {self.XLSX_MAX_ROWS_PER_SHEET} 行"
                    )
                if int(worksheet.max_column or 0) > self.XLSX_MAX_COLUMNS:
                    truncation_notes.append(
                        f"仅解析前 {self.XLSX_MAX_COLUMNS} 列"
                    )

                section_lines = [f"[工作表] {worksheet.title}"]
                section_lines.extend(f"[行 {row_no}] {line}" for row_no, line in rows)
                if truncation_notes:
                    section_lines.append("[解析提示] " + "；".join(truncation_notes))
                document_sections.append("\n".join(section_lines))

                for piece, row_start, row_end in self._chunk_xlsx_rows(
                    worksheet.title,
                    rows,
                    truncation_notes,
                ):
                    chunks.append(
                        ParsedChunk(
                            index=chunk_index,
                            text=piece,
                            sheet_name=worksheet.title,
                            row_start=row_start,
                            row_end=row_end,
                        )
                    )
                    chunk_index += 1

            if len(workbook.worksheets) > self.XLSX_MAX_SHEETS:
                document_sections.append(
                    f"[解析提示] 工作簿共 {len(workbook.worksheets)} 个工作表，"
                    f"仅解析前 {self.XLSX_MAX_SHEETS} 个。"
                )
        finally:
            workbook.close()

        text = "\n\n".join(document_sections).strip()
        if not text or not chunks:
            raise EmptyDocumentError("XLSX 未提取到可读单元格数据")

        return ParsedDocument(
            filename=filename,
            media_type=media_type,
            parser="openpyxl",
            text=text,
            chunks=tuple(chunks),
            page_count=None,
        )

    def _chunk_xlsx_rows(
        self,
        sheet_name: str,
        rows: list[tuple[int, str]],
        truncation_notes: list[str],
    ) -> list[tuple[str, int, int]]:
        result: list[tuple[str, int, int]] = []
        buffer: list[tuple[int, str]] = []
        buffer_len = 0

        def flush() -> None:
            nonlocal buffer, buffer_len
            if not buffer:
                return
            row_start = buffer[0][0]
            row_end = buffer[-1][0]
            header = f"[工作表] {sheet_name}；[行范围] {row_start}-{row_end}"
            body = "\n".join(f"[行 {row_no}] {line}" for row_no, line in buffer)
            note = ""
            if truncation_notes:
                note = "\n[解析提示] " + "；".join(truncation_notes)
            result.append((f"{header}\n{body}{note}", row_start, row_end))
            buffer = []
            buffer_len = 0

        for row_no, line in rows:
            rendered = f"[行 {row_no}] {line}"
            extra = len(rendered) + (1 if buffer else 0)
            if buffer and buffer_len + extra > self.chunk_chars:
                flush()
            buffer.append((row_no, line))
            buffer_len += extra
        flush()
        return result

    @staticmethod
    def _xlsx_cell_text(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.isoformat(sep=" ")
        if isinstance(value, (date, time)):
            return value.isoformat()
        if isinstance(value, bool):
            return "TRUE" if value else "FALSE"
        text = str(value).strip()
        return text.replace("\t", " ").replace("\r\n", " ↵ ").replace("\n", " ↵ ")

    def _chunk_text(self, text: str) -> list[str]:
        cleaned = text.strip()
        if len(cleaned) <= self.chunk_chars:
            return [cleaned] if cleaned else []

        result: list[str] = []
        start = 0
        while start < len(cleaned):
            end = min(len(cleaned), start + self.chunk_chars)
            if end < len(cleaned):
                # Prefer a sentence/line boundary near the end of the window.
                candidates = [
                    cleaned.rfind(mark, start + self.chunk_chars // 2, end)
                    for mark in ("\n", "。", "；", ";", ".")
                ]
                boundary = max(candidates)
                if boundary > start:
                    end = boundary + 1
            piece = cleaned[start:end].strip()
            if piece:
                result.append(piece)
            if end >= len(cleaned):
                break
            start = max(start + 1, end - self.overlap_chars)
        return result


# Backward-compatible alias: existing code can keep importing ChatFileParser.
ChatFileParser = UnifiedFileParser
